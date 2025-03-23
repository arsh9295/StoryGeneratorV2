from docx import Document
import os
from pathlib import Path
import logging

logger = logging.getLogger(__name__)
class WritetoDocx:
    def __init__(self):
        self.document = None
        logger.debug("WritetoDocx instance initialized")

    def create_story_folder(self, file_path):
        try:
            logger.debug(f"Creating story folder for path: {file_path}")
            
            # Create directory if it doesn't exist
            os.makedirs(os.path.dirname(file_path), exist_ok=True)
            
            # Create new document or load existing
            if os.path.exists(file_path):
                self.document = Document(file_path)
                logger.debug("Loaded existing document")
            else:
                self.document = Document()
                self.document.save(file_path)
                logger.debug("Created new document")
            
            return file_path
            
        except Exception as e:
            logger.error(f"Error in create_story_folder: {str(e)}", exc_info=True)
            raise Exception(f"Error creating story folder: {str(e)}")

    def write_chapter(self, file_path, chapter_title, content):
        try:
            logger.debug(f"Writing chapter: {chapter_title}")
            
            if self.document is None:
                logger.debug("Document not initialized, loading/creating document")
                self.create_story_folder(file_path)
            
            self.document.add_page_break()
            self.document.add_heading(chapter_title, level=2)
            self.document.add_paragraph(content)
            self.document.save(file_path)
            logger.debug("Chapter written successfully")
            
            return True
            
        except Exception as e:
            logger.error(f"Error in write_chapter: {str(e)}", exc_info=True)
            raise Exception(f"Error writing chapter: {str(e)}")

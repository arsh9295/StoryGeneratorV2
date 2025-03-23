from fastapi import FastAPI
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
import google.generativeai as genai
import json
from docx import Document
import os
from pathlib import Path
import re
import google.generativeai as genai
from PIL import Image, ImageDraw
import io
import base64
from openai import OpenAI

class DeepseekGenerator:
    def __init__(self):
        self.api_key = None

    def set_api_key(self, key: str):
        if not key or not isinstance(key, str):
            raise ValueError("Invalid API key")
        self.api_key = key
        return True

    def generate_story(self, prompt: str) -> str:
        if not self.api_key:
            raise ValueError("API key not set")
        else:
            genai.configure(api_key=self.api_key)
            model = genai.GenerativeModel('gemini-2.0-flash')
            
            response = model.generate_content(prompt)
            if response.text:
                write_to_docx(chapter_title, response.text, story_file_path)
            return response.text        # Implement your Gemini story generation logic here
        return "Generated story based on prompt: " + (prompt or "default prompt")

    def save_story(self, story_data: dict):
        # Add your story saving logic here
        try:
            # Save story to database or file
            return True
        except Exception as e:
            print(f"Error saving story: {e}")
            return False

    def write_to_docx(self, chapter_title: str, content: str, file_path: str):
        if os.path.exists(story_file):
            doc = Document(story_file)
        else:
            doc = Document()
        doc.add_page_break()
        doc.add_heading(f"{heading}", level=2)
        doc.add_paragraph(f"{content}")
        doc.save(story_file)
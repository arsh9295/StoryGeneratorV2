from fastapi import FastAPI
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
import google.generativeai as genai
import json
from docx import Document
import os
from pathlib import Path
import re
import google.generativeai as genai
from PIL import Image, ImageDraw
import io
import base64
from openai import OpenAI

class WritetoDocx:
    """ "
    "This class handles the writing of content to a .docx file.
    "It checks if the file exists and appends content to it."
    """
    def __init__(self):
        self.api_key = None

    def sanitize_filename(self, filename):
        # Remove invalid characters and replace with underscores
        # Keep some common punctuation that's allowed
        sanitized = re.sub(r'[<>:"/\\|?*]', '_', filename)
        # Remove any leading/trailing periods or spaces
        sanitized = sanitized.strip('. ')
        # Ensure filename is not empty
        if not sanitized:
            sanitized = 'unnamed_story'
        return sanitized

    def create_story_folder(self, filename):
        # base_path = Path("stories")
        # story_dir = base_path / story_language / story_type
        story_dir = Path(os.path.dirname(filename))
        story_dir.mkdir(parents=True, exist_ok=True)
        file = os.path.basename(filename)

        # Sanitize the filename before creating the file
        safe_filename = self.sanitize_filename(file)

        story_file = story_dir / f"{safe_filename}"
        
        return story_file

    def write_to_docx(self, title: str, content: str, file_path: str):
        if os.path.exists(file_path):
            doc = Document(file_path)
        else:
            doc = Document()
        doc.add_page_break()
        doc.add_heading(f"{title}", level=2)
        doc.add_paragraph(f"{content}")
        doc.save(file_path)
        return True